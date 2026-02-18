"""Shared utilities: file reading and PO numbering."""

import re
import io
import base64
import pdfplumber
from docx import Document
from openpyxl import load_workbook


def get_next_po_number(po_files: list[str]) -> str:
    """Return next 3-digit PO number based on existing PO filenames."""
    numbers = []
    for fname in po_files:
        match = re.match(r"(\d{3})", fname)
        if match:
            numbers.append(int(match.group(1)))
    return "001" if not numbers else f"{max(numbers) + 1:03d}"


def read_pdf_text(path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)
    return chr(10).join(parts)


def read_pdf_as_images(path: str) -> list[dict]:
    """Convert each PDF page to a base64 PNG dict for Claude Vision."""
    images = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            buf = io.BytesIO()
            page.to_image(resolution=200).original.save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            images.append({
                "type": "image",
                "source": {"type": "base64", "media_type": "image/png", "data": b64},
            })
    return images


def read_docx_text(path: str) -> str:
    """Extract text from a Word document."""
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = chr(9).join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return chr(10).join(parts)


def read_excel_text(path: str) -> str:
    """Extract text from an Excel workbook."""
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            row_str = chr(9).join(str(c) if c is not None else "" for c in row)
            if row_str.strip():
                parts.append(row_str)
    return chr(10).join(parts)


def read_file_text(path: str) -> str:
    """Read quotation file text regardless of format."""
    ext = path.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        return read_pdf_text(path)
    elif ext == "docx":
        return read_docx_text(path)
    elif ext in ("xlsx", "xls"):
        return read_excel_text(path)
    raise ValueError(f"Unsupported file type: .{ext}")
