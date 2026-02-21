"""
Analyzes an example PO PDF via LLM and produces:
  - po_template.html (Jinja2 HTML) OR po_template.docx
  - po_fields.json (field manifest)
"""

import re
import io
import json
import base64
import pdfplumber
import docx
from pdf2docx import Converter
from core.llm_provider import LLMProvider
from core.utils import read_pdf_as_images_base64
from core.config import DEFAULT_TEMPLATE_PATH, DEFAULT_FIELDS_PATH, DEFAULT_DOCX_TEMPLATE_PATH

DOCX_ANALYZER_PROMPT = """You are creating a Word document template for a Purchase Order.

Identify exactly 9 dynamic fields from the document:
  • po_date, po_number, supplier_company, supplier_contact, subject, 
    delivery_time, payment_term, delivery_term, total_price

For each field, find the EXACT literal text current value in the document and map it to its placeholder.
Return a JSON object with a "mappings" key.

Example:
{
  "mappings": {
    "03.01.2026": "{{ po_date }}",
    "548": "{{ po_number }}",
    "Acme Corp": "{{ supplier_company }}"
  }
}

IMPORTANT: Return ONLY the JSON object. No conversational text before or after.
"""

HTML_ANALYZER_PROMPT = """You are a professional document engineer. Your task is to create a high-fidelity HTML/CSS template that EXACTLY replicates the provided Purchase Order image.

FIELD CLASSIFICATION RULES:
1. STATIC DATA (Hardcode these in HTML):
   - Buyer Company Name (e.g., Güney Yıldızı Petrol, GYP)
   - Buyer Addresses and Fixed Legal Text.
   - Fixed Signature titles or company names.
   - LOGOS: Recreate the LOGO visually using styled <div> elements and text. For example, if there is a 'GYP' box, create it with CSS background colors and borders. 
   - DO NOT create {{ placeholders }} for these.
   - IMPORTANT: DO NOT add any extra text, icons, or asterisks (*) that are not in the original document image.

2. DYNAMIC DATA (Use {{ field_name }} placeholders):
   - Order-specific data only: po_date, po_number, subject.
   - Vendor/Supplier info: supplier_company, supplier_contact/attn.
   - Terms: delivery_time, payment_term, delivery_term.
   - Financials: total_price (including currency).

LAYOUT RULES:
1. Use HTML TABLES (<table>) for the primary layout structure to ensure 100% stability.
2. MATCH ORIGINAL DENSITY: Replicate the line-height, font-size, and spacing of the original document. 
   - If the original is dense, use tight `line-height: 1.1;` and minimal `padding: 1px 3px;`.
   - Ensure the generated PO does not exceed the page count of the original unless the content requires it.
3. COMPACTNESS: Avoid adding default HTML margins or extra vertical whitespace that is not present in the original document image.
4. Set the main container to 210mm width (A4) with standard margins (e.g., 10-15mm).
5. Use standard web-safe fonts (Arial, Helvetica, sans-serif).
6. IMPORTANT: DO NOT add any extra text, icons, or asterisks (*) that are not in the original document.
7. LOGOS: Do not try to draw logos with symbols. Use a <div> with the text "[LOGO]" or leave space.

OUTPUT FORMAT:
Return your response in this EXACT structure (DO NOT SKIP ANY TAGS):

<html_template>
[complete HTML/CSS here]
</html_template>
<field_manifest>
{
  "fields": [
    {"name": "field_name", "description": "Short description"},
    ...
  ]
}
</field_manifest>
"""

def parse_analyzer_response(response_text: str) -> tuple[str, list[dict]]:
    """Parse structured response into (html_str, fields_list)."""
    # 1. Extract HTML
    html_match = re.search(r"<html_template>([\s\S]*?)</html_template>", response_text)
    if html_match:
        html = html_match.group(1).strip()
    else:
        # Fallback: find anything that looks like HTML
        if "<!DOCTYPE html>" in response_text:
            html = response_text.split("<!DOCTYPE html>")[1].split("</html_template>")[0]
            html = "<!DOCTYPE html>" + html
        elif "<html>" in response_text:
            html = response_text.split("<html>")[1].split("</html>")[0]
            html = "<html>" + html + "</html>"
        else:
            raise ValueError("Response missing <html_template> block and no HTML found.")

    # 2. Extract Manifest
    manifest_match = re.search(r"<field_manifest>([\s\S]*?)</field_manifest>", response_text)
    if manifest_match:
        manifest_json = manifest_match.group(1).strip()
    else:
        # Fallback: find any JSON block that contains "fields"
        json_match = re.search(r"\{[\s\S]*?\"fields\"[\s\S]*?\}", response_text)
        if json_match:
            manifest_json = json_match.group()
        else:
            raise ValueError("Response missing <field_manifest> block and no JSON manifest found.")

    try:
        manifest = json.loads(manifest_json)
        return html, manifest["fields"]
    except Exception as e:
        raise ValueError(f"Failed to parse manifest JSON: {str(e)}")

def analyze_example_po_to_html(
    pdf_path: str,
    config: dict,
    template_path: str = DEFAULT_TEMPLATE_PATH,
    fields_path: str = DEFAULT_FIELDS_PATH,
) -> list[dict]:
    """
    Analyzes only the FIRST page of the PO PDF and produces a high-fidelity HTML template.
    """
    provider = LLMProvider(
        provider=config.get("provider", "anthropic"),
        api_key=config["api_key"],
        model=config["model"]
    )
    
    all_images = read_pdf_as_images_base64(pdf_path)
    if not all_images:
        raise ValueError("Could not read images from PDF.")
    
    # Strictly use the first page to ignore attached quotations
    first_page_image = [all_images[0]]
    response = provider.generate_with_vision(HTML_ANALYZER_PROMPT, first_page_image)
    
    try:
        html, fields = parse_analyzer_response(response)
    except Exception as e:
        # Save raw response for debugging
        with open("debug_response.txt", "w", encoding="utf-8") as f:
            f.write(response)
        raise RuntimeError(f"Şablon ayrıştırılamadı: {str(e)}. Lütfen debug_response.txt dosyasını kontrol edin.")

    with open(template_path, "w", encoding="utf-8") as f:
        f.write(html)
    with open(fields_path, "w", encoding="utf-8") as f:
        json.dump({"fields": fields}, f, indent=2, ensure_ascii=False)

    return fields

def pdf_to_docx(pdf_path: str, docx_path: str):
    """Convert only the first page of the PDF to DOCX to avoid including quotation pages."""
    cv = Converter(pdf_path)
    # start=0, end=1 means process only the first page
    cv.convert(docx_path, start=0, end=1)
    cv.close()

def _replace_text_in_docx(docx_path: str, mapping: dict):
    doc = docx.Document(docx_path)
    # Sort keys by length (longest first) to avoid partial replacements
    sorted_keys = sorted(mapping.keys(), key=len, reverse=True)
    
    for old_text in sorted_keys:
        new_placeholder = mapping[old_text]
        if not old_text or len(str(old_text)) < 2: continue
        
        for p in doc.paragraphs:
            if old_text in p.text: p.text = p.text.replace(old_text, new_placeholder)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text: p.text = p.text.replace(old_text, new_placeholder)
    doc.save(docx_path)

def analyze_example_po_to_docx(
    pdf_path: str,
    config: dict,
    docx_path: str = DEFAULT_DOCX_TEMPLATE_PATH,
    fields_path: str = DEFAULT_FIELDS_PATH,
) -> list[dict]:
    # 1. Convert ONLY the first page to DOCX
    pdf_to_docx(pdf_path, docx_path)
    
    provider = LLMProvider(
        provider=config.get("provider", "anthropic"),
        api_key=config["api_key"],
        model=config["model"]
    )
    
    # 2. Only send the first page image to the AI for analysis
    all_images = read_pdf_as_images_base64(pdf_path)
    first_page_image = [all_images[0]] if all_images else []
    
    response = provider.generate_with_vision(DOCX_ANALYZER_PROMPT, first_page_image)
    
    # EVEN MORE ROBUST JSON PARSING
    try:
        # 1. Try finding JSON object directly
        json_match = re.search(r"\{[\s\S]*\}", response)
        if not json_match:
            raise ValueError("JSON block not found in response.")
        
        json_str = json_match.group()
        # Clean potential markdown or extra characters
        json_str = json_str.strip().strip("`").replace("json\n", "", 1)
        data = json.loads(json_str)
        
        # 2. Extract mapping
        if "mappings" in data:
            mapping = data["mappings"]
        elif isinstance(data, dict):
            # Check if any value has Jinja brackets
            if any("{{" in str(v) for v in data.values()):
                mapping = data
            else:
                # Last resort fallback if structure is weird
                mapping = data
        else:
            raise ValueError("Structure of JSON response is not a dictionary.")

        _replace_text_in_docx(docx_path, mapping)
        
    except Exception as e:
        # Log the full response for debugging if needed
        with open("debug_response.txt", "w", encoding="utf-8") as f:
            f.write(response)
        raise RuntimeError(f"JSON Ayrıştırma Hatası: {str(e)}. Lütfen debug_response.txt dosyasını kontrol edin.")

    fields = [
        {"name": "po_date", "description": "Sipariş Tarihi"},
        {"name": "po_number", "description": "PO Numarası"},
        {"name": "supplier_company", "description": "Tedarikçi Firma"},
        {"name": "supplier_contact", "description": "İlgili Kişi"},
        {"name": "subject", "description": "Konu"},
        {"name": "delivery_time", "description": "Teslim Süresi"},
        {"name": "payment_term", "description": "Ödeme Şekli"},
        {"name": "delivery_term", "description": "Teslim Şekli"},
        {"name": "total_price", "description": "Toplam Tutar"},
    ]
    with open(fields_path, "w", encoding="utf-8") as f:
        json.dump({"fields": fields}, f, indent=2, ensure_ascii=False)
    return fields
