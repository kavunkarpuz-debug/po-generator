"""
PO (Purchase Order) Generator
Reads a supplier quotation, extracts data via Claude API,
generates a PO cover letter from template, and merges into a final PDF.
"""

import os
import sys
import re
import json
import base64
import datetime

import anthropic
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from pypdf import PdfWriter, PdfReader
from docx2pdf import convert as docx2pdf_convert

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
WORKING_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_NAME = "PO_TEMPLATE.docx"
SCRIPT_NAME = "po_generator.py"
SKIP_NAMES = {TEMPLATE_NAME.lower(), SCRIPT_NAME.lower(),
              "po_automation_instructions.md", "po_kurulum_rehberi.md"}
VALID_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls"}


# ---------------------------------------------------------------------------
# 1. File identification
# ---------------------------------------------------------------------------
def find_quotation_and_pos(directory: str):
    """Return (quotation_path, list_of_po_paths)."""
    quotation_files = []
    po_files = []

    for fname in os.listdir(directory):
        fpath = os.path.join(directory, fname)
        if not os.path.isfile(fpath):
            continue
        name_lower = fname.lower()
        ext = os.path.splitext(name_lower)[1]
        if ext not in VALID_EXTENSIONS:
            continue
        if name_lower in SKIP_NAMES:
            continue
        if name_lower.endswith(".md"):
            continue

        if "po" in name_lower:
            po_files.append(fname)
        else:
            quotation_files.append(fname)

    return quotation_files, po_files


# ---------------------------------------------------------------------------
# 2. PO numbering
# ---------------------------------------------------------------------------
def get_next_po_number(po_files: list[str]) -> str:
    """Extract highest 3-digit prefix from PO filenames, return next number."""
    numbers = []
    for fname in po_files:
        match = re.match(r"(\d{3})", fname)
        if match:
            numbers.append(int(match.group(1)))

    if not numbers:
        return "001"
    return f"{max(numbers) + 1:03d}"


# ---------------------------------------------------------------------------
# 3. Read quotation content
# ---------------------------------------------------------------------------
def read_pdf_text(path: str) -> str:
    """Extract text from PDF using pdfplumber."""
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def read_pdf_as_images(path: str) -> list[dict]:
    """Convert PDF pages to base64 images for Claude Vision API."""
    images = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            img = page.to_image(resolution=200)
            import io
            buf = io.BytesIO()
            img.original.save(buf, format="PNG")
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            images.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": b64,
                },
            })
    return images


def read_docx_text(path: str) -> str:
    """Extract text from Word document."""
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)


def read_excel_text(path: str) -> str:
    """Extract text from Excel workbook."""
    wb = load_workbook(path, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"--- Sheet: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            row_str = "\t".join(str(c) if c is not None else "" for c in row)
            if row_str.strip():
                parts.append(row_str)
    return "\n".join(parts)


def read_quotation(path: str) -> str:
    """Read quotation file based on extension. Returns text content."""
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return read_pdf_text(path)
    elif ext == ".docx":
        return read_docx_text(path)
    elif ext in (".xlsx", ".xls"):
        return read_excel_text(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------------------------------------------------------
# 4. Claude API extraction
# ---------------------------------------------------------------------------
EXTRACTION_PROMPT = """You are a document parser for purchase orders. Extract the following fields from this supplier quotation document.

Return ONLY a valid JSON object with these exact keys:
- supplier_company: Full company name of the supplier/seller
- supplier_contact: Name of the contact person/issuer at the supplier
- subject: Brief description of what is being purchased (3-5 words MAX, like 'BOP Rings' or 'Mud Pump Liners')
- delivery_time: Delivery time/period as stated
- payment_term: Payment terms as stated
- delivery_term: Delivery/shipping terms
- total_price: Total price with currency (the grand total, all-inclusive)

If a field cannot be determined, use "N/A".

DOCUMENT CONTENT:
{document_text}"""

EXTRACTION_PROMPT_VISION = """You are a document parser for purchase orders. Look at this supplier quotation document image(s) and extract the following fields.

Return ONLY a valid JSON object with these exact keys:
- supplier_company: Full company name of the supplier/seller
- supplier_contact: Name of the contact person/issuer at the supplier
- subject: Brief description of what is being purchased (3-5 words MAX, like 'BOP Rings' or 'Mud Pump Liners')
- delivery_time: Delivery time/period as stated
- payment_term: Payment terms as stated
- delivery_term: Delivery/shipping terms
- total_price: Total price with currency (the grand total, all-inclusive)

If a field cannot be determined, use "N/A"."""


def extract_data_with_claude(quotation_path: str) -> dict:
    """Call Claude API to extract structured data from quotation."""
    client = anthropic.Anthropic()

    # First try text extraction
    text = read_quotation(quotation_path)

    if text and len(text.strip()) > 50:
        print("  Extracting data via Claude API (text mode)...")
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            messages=[{
                "role": "user",
                "content": EXTRACTION_PROMPT.format(document_text=text),
            }],
        )
        response_text = message.content[0].text
    else:
        # Fall back to vision if text extraction yielded little
        ext = os.path.splitext(quotation_path)[1].lower()
        if ext != ".pdf":
            raise RuntimeError(
                f"Text extraction failed and vision fallback only supports PDF, got {ext}"
            )
        print("  Text extraction insufficient, using Claude Vision API...")
        images = read_pdf_as_images(quotation_path)
        content = images + [{"type": "text", "text": EXTRACTION_PROMPT_VISION}]
        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1024,
            messages=[{"role": "user", "content": content}],
        )
        response_text = message.content[0].text

    # Parse JSON from response
    try:
        # Try to extract JSON from possible markdown code block
        json_match = re.search(r"\{[\s\S]*\}", response_text)
        if json_match:
            data = json.loads(json_match.group())
        else:
            data = json.loads(response_text)
    except json.JSONDecodeError:
        print(f"  WARNING: Failed to parse Claude response as JSON.")
        print(f"  Response was: {response_text[:500]}")
        raise

    # Ensure all required keys exist
    required_keys = [
        "supplier_company", "supplier_contact", "subject",
        "delivery_time", "payment_term", "delivery_term", "total_price",
    ]
    for key in required_keys:
        if key not in data:
            data[key] = "N/A"

    return data


# ---------------------------------------------------------------------------
# 5. Word document generation (placeholder replacement)
# ---------------------------------------------------------------------------
def replace_placeholder_in_runs(paragraph, placeholder: str, value: str):
    """Replace a placeholder in a paragraph, handling split-across-runs case."""
    # Fast path: placeholder is within a single run
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            return True

    # Slow path: placeholder may be split across multiple runs
    full_text = paragraph.text
    if placeholder not in full_text:
        return False

    # Reconstruct: find which runs contain parts of the placeholder
    runs = paragraph.runs
    if not runs:
        return False

    # Build a mapping of character positions to runs
    char_positions = []  # (run_index, char_index_in_run)
    for ri, run in enumerate(runs):
        for ci in range(len(run.text)):
            char_positions.append((ri, ci))

    # Find placeholder start in full text
    start = full_text.index(placeholder)
    end = start + len(placeholder)

    # Determine which runs are affected
    start_run, start_char = char_positions[start]
    end_run, end_char = char_positions[end - 1]

    # Put replacement text in the first affected run at the right position
    first_run_text = runs[start_run].text
    runs[start_run].text = first_run_text[:start_char] + value + runs[end_run].text[end_char + 1:]

    # Clear text from intermediate and end runs
    for ri in range(start_run + 1, end_run + 1):
        runs[ri].text = ""

    return True


def generate_po_docx(template_path: str, output_path: str, replacements: dict[str, str]):
    """Copy template and replace all placeholders."""
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            replace_placeholder_in_runs(paragraph, placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, value in replacements.items():
                        replace_placeholder_in_runs(paragraph, placeholder, value)

    doc.save(output_path)


# ---------------------------------------------------------------------------
# 6. PDF conversion and merging
# ---------------------------------------------------------------------------
def convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """Convert Word document to PDF using docx2pdf (MS Word COM on Windows)."""
    print(f"  Converting {os.path.basename(docx_path)} to PDF...")
    docx2pdf_convert(docx_path, pdf_path)


def convert_quotation_to_pdf(quotation_path: str, working_dir: str) -> str:
    """Convert non-PDF quotation to PDF. Returns path to the PDF."""
    ext = os.path.splitext(quotation_path)[1].lower()
    if ext == ".pdf":
        return quotation_path

    base = os.path.splitext(os.path.basename(quotation_path))[0]
    temp_pdf = os.path.join(working_dir, f"_temp_{base}.pdf")

    if ext == ".docx":
        docx2pdf_convert(quotation_path, temp_pdf)
    elif ext in (".xlsx", ".xls"):
        # For Excel, convert via docx2pdf won't work; use a workaround
        # Create a temporary docx with the excel content, then convert
        # Actually, docx2pdf only works with .docx. For Excel we need another approach.
        # We'll try using win32com directly for Excel.
        try:
            import win32com.client
            excel = win32com.client.Dispatch("Excel.Application")
            excel.Visible = False
            wb = excel.Workbooks.Open(os.path.abspath(quotation_path))
            wb.ExportAsFixedFormat(0, os.path.abspath(temp_pdf))
            wb.Close(False)
            excel.Quit()
        except Exception as e:
            print(f"  WARNING: Could not convert Excel to PDF: {e}")
            print("  The merged PDF will only contain the PO cover letter.")
            return None
    else:
        print(f"  WARNING: Cannot convert {ext} to PDF.")
        return None

    return temp_pdf


def merge_pdfs(po_pdf_path: str, quotation_pdf_path: str, output_path: str):
    """Merge PO cover letter PDF with quotation PDF."""
    writer = PdfWriter()

    po_reader = PdfReader(po_pdf_path)
    for page in po_reader.pages:
        writer.add_page(page)

    if quotation_pdf_path and os.path.exists(quotation_pdf_path):
        quote_reader = PdfReader(quotation_pdf_path)
        for page in quote_reader.pages:
            writer.add_page(page)

    with open(output_path, "wb") as f:
        writer.write(f)


# ---------------------------------------------------------------------------
# 7. Main workflow
# ---------------------------------------------------------------------------
def main():
    print("=" * 60)
    print("  PO Generator - Purchase Order Automation")
    print("=" * 60)
    print(f"\nWorking directory: {WORKING_DIR}")

    # Step 1: Identify files
    print("\n[1/7] Scanning directory...")
    quotation_files, po_files = find_quotation_and_pos(WORKING_DIR)

    if not quotation_files:
        print("ERROR: No quotation file found in the directory.")
        print("Place a quotation file (PDF/DOCX/XLSX) without 'PO' in its name.")
        sys.exit(1)

    if len(quotation_files) > 1:
        print("ERROR: Multiple quotation files found:")
        for f in quotation_files:
            print(f"  - {f}")
        print("Please keep only ONE quotation file in the directory.")
        sys.exit(1)

    quotation_file = quotation_files[0]
    quotation_path = os.path.join(WORKING_DIR, quotation_file)
    print(f"  Quotation: {quotation_file}")
    print(f"  Existing PO files: {len(po_files)}")

    # Step 2: Determine PO number
    print("\n[2/7] Determining PO number...")
    po_number = get_next_po_number(po_files)
    print(f"  Next PO number: {po_number}")

    # Step 3: Extract data from quotation
    print("\n[3/7] Reading and extracting data from quotation...")
    try:
        extracted = extract_data_with_claude(quotation_path)
    except Exception as e:
        print(f"  First attempt failed: {e}")
        print("  Retrying once...")
        try:
            extracted = extract_data_with_claude(quotation_path)
        except Exception as e2:
            print(f"  Retry also failed: {e2}")
            print("\nPlease enter data manually:")
            extracted = {}
            fields = [
                ("supplier_company", "Supplier company name"),
                ("supplier_contact", "Contact person name"),
                ("subject", "Subject (3-5 words, e.g. 'BOP Rings')"),
                ("delivery_time", "Delivery time"),
                ("payment_term", "Payment terms"),
                ("delivery_term", "Delivery terms"),
                ("total_price", "Total price with currency"),
            ]
            for key, prompt in fields:
                val = input(f"  {prompt}: ").strip()
                extracted[key] = val if val else "N/A"

    print("\n  Extracted data:")
    for k, v in extracted.items():
        print(f"    {k}: {v}")

    # Step 4: Build filename and replacements
    print("\n[4/7] Preparing document...")
    today = datetime.date.today()
    date_ddmmyyyy = today.strftime("%d%m%Y")
    date_dotted = today.strftime("%d.%m.%Y")
    subject = extracted.get("subject", "N/A")
    po_no_full = f"{po_number}_{date_ddmmyyyy}"
    base_filename = f"{po_number}_{date_ddmmyyyy} PO for {subject}"

    docx_output = os.path.join(WORKING_DIR, f"{base_filename}.docx")
    final_pdf_output = os.path.join(WORKING_DIR, f"{base_filename}.pdf")

    replacements = {
        "{{DATE}}": date_dotted,
        "{{SUPPLIER}}": extracted.get("supplier_company", "N/A"),
        "{{ATTN}}": extracted.get("supplier_contact", "N/A"),
        "{{PO_NO}}": po_no_full,
        "{{SUBJECT}}": subject,
        "{{DELIVERY_TIME}}": extracted.get("delivery_time", "N/A"),
        "{{PAYMENT_TERM}}": extracted.get("payment_term", "N/A"),
        "{{DELIVERY_TERM}}": extracted.get("delivery_term", "N/A"),
        "{{TOTAL_PRICE}}": extracted.get("total_price", "N/A"),
    }

    print(f"  Filename: {base_filename}")
    print(f"  PO No: {po_no_full}")

    # Step 5: Generate Word document
    print("\n[5/7] Generating PO Word document...")
    template_path = os.path.join(WORKING_DIR, TEMPLATE_NAME)
    if not os.path.exists(template_path):
        print(f"ERROR: Template file not found: {template_path}")
        sys.exit(1)

    generate_po_docx(template_path, docx_output, replacements)
    print(f"  Created: {os.path.basename(docx_output)}")

    # Step 6: Convert Word to PDF
    print("\n[6/7] Converting to PDF...")
    temp_po_pdf = os.path.join(WORKING_DIR, f"_temp_po_{po_number}.pdf")
    try:
        convert_docx_to_pdf(docx_output, temp_po_pdf)
    except Exception as e:
        print(f"  ERROR converting Word to PDF: {e}")
        print("  Make sure Microsoft Word is installed.")
        print(f"  The .docx file was still created: {os.path.basename(docx_output)}")
        sys.exit(1)

    # Step 7: Merge PDFs
    print("\n[7/7] Merging PO cover letter with quotation...")
    quotation_pdf_path = convert_quotation_to_pdf(quotation_path, WORKING_DIR)
    merge_pdfs(temp_po_pdf, quotation_pdf_path, final_pdf_output)
    print(f"  Created: {os.path.basename(final_pdf_output)}")

    # Cleanup temp files
    if os.path.exists(temp_po_pdf):
        os.remove(temp_po_pdf)
    if quotation_pdf_path and quotation_pdf_path != quotation_path and os.path.exists(quotation_pdf_path):
        os.remove(quotation_pdf_path)

    # Summary
    print("\n" + "=" * 60)
    print("  DONE! New files created:")
    print(f"  1. {os.path.basename(docx_output)}")
    print(f"  2. {os.path.basename(final_pdf_output)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
